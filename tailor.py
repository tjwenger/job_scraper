"""
Resume tailoring — uses Claude Sonnet to rewrite resume.txt for a specific job.
Returns plain-text resume in the same structured format as the original.
"""
import os
import re
import io
from pathlib import Path
from anthropic import Anthropic

_client = None


def _get_client() -> Anthropic:
    global _client
    if _client is None:
        _client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    return _client


def _load_resume() -> str:
    path = Path(__file__).parent / "resume.txt"
    return path.read_text(encoding="utf-8")


def tailor_resume(title: str, company: str, description: str) -> str:
    """Call Claude Sonnet to tailor the resume for the given job. Returns plain text."""
    resume = _load_resume()
    client = _get_client()

    prompt = f"""You are an expert resume writer. Tailor the resume below for the specific job posting provided.

RULES:
- Keep all facts truthful — do not invent experience, skills, or dates
- Rewrite the SUMMARY section to speak directly to this role and company
- Reorder and emphasize bullet points in EXPERIENCE that are most relevant to this job
- Mirror keywords and language from the job description naturally (no keyword stuffing)
- Keep the same sections and plain-text format as the original
- Do not add new jobs, change companies, or alter dates
- Return ONLY the full tailored resume text — no commentary, no markdown fences

JOB TITLE: {title}
COMPANY: {company}

JOB DESCRIPTION:
{description or "(No description available — tailor based on title and company)"}

ORIGINAL RESUME:
{resume}"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text.strip()


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def _parse_resume(text: str) -> dict:
    """
    Parse the plain-text resume into a dict of sections.
    Returns: {
        "header": ["Name line", "Contact line"],
        "sections": [{"title": "SUMMARY", "lines": [...], "entries": [...]}]
    }
    Each entry in EXPERIENCE has: {"heading": str, "bullets": [str]}
    """
    lines = [l.rstrip() for l in text.splitlines()]

    # First non-empty lines before the first ALL-CAPS section are the header
    header = []
    section_start = 0
    for i, line in enumerate(lines):
        if line and line.isupper() and len(line) > 2:
            section_start = i
            break
        if line:
            header.append(line)

    sections = []
    current_section = None

    for line in lines[section_start:]:
        # Detect section header: all-caps, no leading whitespace, 3+ chars
        if line and line.isupper() and len(line) > 2 and not line.startswith("-"):
            if current_section:
                sections.append(current_section)
            current_section = {"title": line, "lines": [], "entries": []}
        elif current_section is not None:
            current_section["lines"].append(line)

    if current_section:
        sections.append(current_section)

    # Parse EXPERIENCE entries (lines matching "Title — Company | Dates")
    for section in sections:
        if "EXPERIENCE" in section["title"] or "WORK" in section["title"]:
            entry = None
            for line in section["lines"]:
                if not line:
                    continue
                if line.startswith("-") or line.startswith("•"):
                    if entry:
                        entry["bullets"].append(line.lstrip("-• ").strip())
                else:
                    # Looks like a job heading
                    if entry:
                        section["entries"].append(entry)
                    entry = {"heading": line, "bullets": []}
            if entry:
                section["entries"].append(entry)

    return {"header": header, "sections": sections}


def generate_docx(tailored_text: str) -> bytes:
    """Convert tailored resume plain text to a formatted DOCX. Returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    parsed = _parse_resume(tailored_text)

    DARK   = RGBColor(0x1a, 0x1a, 0x2e)
    ACCENT = RGBColor(0x2d, 0x6a, 0xd4)
    GRAY   = RGBColor(0x55, 0x55, 0x55)

    def _set_font(run, size, bold=False, color=None):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def _para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        return p, p.add_run(text) if text else (p, None)

    # --- Header ---
    if parsed["header"]:
        # Name line — bold, large
        name_line = parsed["header"][0]
        # Split on " — " to separate name from title if present
        parts = name_line.split(" — ", 1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(parts[0])
        _set_font(r, 20, bold=True, color=DARK)
        if len(parts) > 1:
            r2 = p.add_run(f"  |  {parts[1]}")
            _set_font(r2, 11, color=GRAY)

        for contact_line in parsed["header"][1:]:
            p2 = doc.add_paragraph(contact_line)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(0)
            for run in p2.runs:
                _set_font(run, 10, color=GRAY)

    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    border_run = p.add_run()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2D6AD4")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Sections ---
    for section in parsed["sections"]:
        title = section["title"]

        # Section header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(title)
        _set_font(r, 11, bold=True, color=ACCENT)

        # Thin underline under section title
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

        if section["entries"]:
            # EXPERIENCE-style section
            for entry in section["entries"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(1)
                # Bold the job title part, regular for company/dates
                heading = entry["heading"]
                # Split on " — " or " | " to bold just the title
                m = re.match(r"^([^—|]+?)(?:\s[—|]\s|\s—\s|\s\|\s)(.+)$", heading)
                if m:
                    r1 = p.add_run(m.group(1).strip())
                    _set_font(r1, 10, bold=True, color=DARK)
                    r2 = p.add_run(f"  —  {m.group(2).strip()}")
                    _set_font(r2, 10, color=GRAY)
                else:
                    r1 = p.add_run(heading)
                    _set_font(r1, 10, bold=True, color=DARK)

                for bullet in entry["bullets"]:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent  = Inches(0.25)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after  = Pt(1)
                    r = p.add_run(bullet)
                    _set_font(r, 10, color=DARK)
        else:
            # Plain text section (SUMMARY, SKILLS, etc.)
            body = "\n".join(l for l in section["lines"] if l).strip()
            if body:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(3)
                r = p.add_run(body)
                _set_font(r, 10, color=DARK)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_pdf(tailored_text: str) -> bytes:
    """Convert tailored resume plain text to a formatted PDF. Returns bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    DARK   = colors.HexColor("#1a1a2e")
    ACCENT = colors.HexColor("#2d6ad4")
    GRAY   = colors.HexColor("#555555")

    styles = {
        "name":    ParagraphStyle("name",    fontName="Helvetica-Bold", fontSize=18, textColor=DARK,   alignment=TA_CENTER, spaceAfter=2),
        "contact": ParagraphStyle("contact", fontName="Helvetica",      fontSize=10, textColor=GRAY,   alignment=TA_CENTER, spaceAfter=4),
        "section": ParagraphStyle("section", fontName="Helvetica-Bold", fontSize=11, textColor=ACCENT, spaceBefore=12, spaceAfter=3),
        "job":     ParagraphStyle("job",     fontName="Helvetica-Bold", fontSize=10, textColor=DARK,   spaceBefore=6, spaceAfter=1),
        "company": ParagraphStyle("company", fontName="Helvetica",      fontSize=10, textColor=GRAY),
        "bullet":  ParagraphStyle("bullet",  fontName="Helvetica",      fontSize=10, textColor=DARK,   leftIndent=14, firstLineIndent=-8, spaceAfter=2),
        "body":    ParagraphStyle("body",    fontName="Helvetica",      fontSize=10, textColor=DARK,   spaceAfter=4),
    }

    parsed = _parse_resume(tailored_text)
    story = []

    # Header
    if parsed["header"]:
        name_line = parsed["header"][0]
        parts = name_line.split(" — ", 1)
        name_html = f"<b>{parts[0]}</b>"
        if len(parts) > 1:
            name_html += f'  <font color="#555555" size="11">|  {parts[1]}</font>'
        story.append(Paragraph(name_html, styles["name"]))
        for line in parsed["header"][1:]:
            story.append(Paragraph(line, styles["contact"]))

    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=6))

    # Sections
    for section in parsed["sections"]:
        story.append(Paragraph(section["title"], styles["section"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))

        if section["entries"]:
            for entry in section["entries"]:
                heading = entry["heading"]
                m = re.match(r"^([^—|]+?)(?:\s[—|]\s|\s—\s|\s\|\s)(.+)$", heading)
                if m:
                    job_html = f"<b>{m.group(1).strip()}</b>"
                    story.append(Paragraph(job_html, styles["job"]))
                    story.append(Paragraph(m.group(2).strip(), styles["company"]))
                else:
                    story.append(Paragraph(f"<b>{heading}</b>", styles["job"]))
                for bullet in entry["bullets"]:
                    story.append(Paragraph(f"• {bullet}", styles["bullet"]))
        else:
            body = " ".join(l for l in section["lines"] if l).strip()
            if body:
                story.append(Paragraph(body, styles["body"]))

    doc.build(story)
    buf.seek(0)
    return buf.read()
